"""
File extraction and redaction for PDF / DOCX inputs.

Public API:
  extract_text(path, suffix)                   -> str
  extract_pdf_with_metadata(path)              -> PdfExtractionResult
  redact_pdf_bytes(path, extraction, spans)    -> bytes
  redact_docx_bytes(path, values)              -> bytes
  UnsupportedFileTypeError
"""

from __future__ import annotations

import io
import shutil
import subprocess
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path


class UnsupportedFileTypeError(ValueError):
    pass


PDF_PAGE_SEPARATOR = "\n\f\n"
PDF_LOW_TEXT_FONT_THRESHOLD = 3.0
PDF_LINE_Y_TOLERANCE = 4.0
PDF_TABLE_GAP_THRESHOLD = 60.0
PDF_WORD_GAP_THRESHOLD = 1.5
PDF_RECT_GAP_THRESHOLD = 6.0
PIPE_TOKENS = {"|", "｜"}


@dataclass(slots=True)
class PdfRectBox:
    page_index: int
    x0: float
    y0: float
    x1: float
    y1: float


@dataclass(slots=True)
class PdfWordBox:
    text: str
    x0: float
    y0: float
    x1: float
    y1: float


@dataclass(slots=True)
class PdfNormalizedPage:
    text: str
    char_boxes: list[PdfRectBox | None]
    provider: str


@dataclass(slots=True)
class PdfExtractionResult:
    text: str
    page_count: int
    used_ocr: bool
    text_layer_char_count: int
    ocr_char_count: int
    quality_flags: list[str]
    pdf_provider: str
    page_providers: list[str]
    char_boxes: list[PdfRectBox | None]

    def rect_boxes_for_span(self, start: int, end: int) -> list[PdfRectBox]:
        boxes: list[PdfRectBox] = []
        for box in self.char_boxes[start:end]:
            if box is not None:
                boxes.append(box)
        return boxes


def _collect_pdf_quality_flags(page: "fitz.Page", text: str) -> set[str]:
    flags: set[str] = set()
    stripped = text.strip()

    if not stripped:
        flags.add("no_text_layer")

    if page.get_images(full=True):
        flags.add("image_heavy")

    if stripped:
        sizes: list[float] = []
        raw = page.get_text("rawdict")
        for block in raw.get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    size = span.get("size")
                    if isinstance(size, (int, float)):
                        sizes.append(float(size))
        if sizes and min(sizes) <= PDF_LOW_TEXT_FONT_THRESHOLD:
            flags.add("low_text_density")

    return flags


def _extract_pdf_with_ocr(path: str, page_numbers: list[int] | None = None) -> list[object]:
    from detectors.pdf_ocr import extract_pdf_pages_with_ocr

    return extract_pdf_pages_with_ocr(path, page_numbers=page_numbers)


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _find_pdftotext_binary() -> str | None:
    return shutil.which("pdftotext")


def _load_pdftotext_word_pages(path: str) -> list[list[PdfWordBox]] | None:
    pdftotext_bin = _find_pdftotext_binary()
    if not pdftotext_bin:
        return None

    result = subprocess.run(
        [pdftotext_bin, "-bbox-layout", path, "-"],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0 or not result.stdout.strip():
        return None

    try:
        root = ET.fromstring(result.stdout)
    except ET.ParseError:
        return None

    pages: list[list[PdfWordBox]] = []
    for page_elem in root.iter():
        if _local_name(page_elem.tag) != "page":
            continue
        words: list[PdfWordBox] = []
        for word_elem in page_elem.iter():
            if _local_name(word_elem.tag) != "word":
                continue
            text = "".join(word_elem.itertext())
            if not text:
                continue
            try:
                words.append(
                    PdfWordBox(
                        text=text,
                        x0=float(word_elem.attrib["xMin"]),
                        y0=float(word_elem.attrib["yMin"]),
                        x1=float(word_elem.attrib["xMax"]),
                        y1=float(word_elem.attrib["yMax"]),
                    )
                )
            except (KeyError, ValueError):
                continue
        pages.append(words)
    return pages


def _append_literal(parts: list[str], char_boxes: list[PdfRectBox | None], text: str) -> None:
    if not text:
        return
    parts.append(text)
    char_boxes.extend([None] * len(text))


def _distribute_word_boxes(word: PdfWordBox, page_index: int) -> list[PdfRectBox | None]:
    if not word.text:
        return []
    width = max(word.x1 - word.x0, 0.001)
    boxes: list[PdfRectBox | None] = []
    length = len(word.text)
    for index, char in enumerate(word.text):
        if char.isspace():
            boxes.append(None)
            continue
        char_x0 = word.x0 + (width * index / length)
        char_x1 = word.x0 + (width * (index + 1) / length)
        boxes.append(PdfRectBox(page_index, char_x0, word.y0, char_x1, word.y1))
    return boxes


def _append_positioned_text(
    parts: list[str],
    char_boxes: list[PdfRectBox | None],
    text: str,
    boxes: list[PdfRectBox | None],
) -> None:
    parts.append(text)
    if len(boxes) == len(text):
        char_boxes.extend(boxes)
        return
    char_boxes.extend([None] * len(text))


def _group_words_into_lines(words: list[PdfWordBox]) -> list[list[PdfWordBox]]:
    lines: list[dict[str, object]] = []
    for word in sorted(words, key=lambda item: (item.y0, item.x0)):
        y_mid = (word.y0 + word.y1) / 2
        target_line: dict[str, object] | None = None
        for line in lines:
            if abs(float(line["y_mid"]) - y_mid) <= PDF_LINE_Y_TOLERANCE:
                target_line = line
                break
        if target_line is None:
            target_line = {"y_mid": y_mid, "words": []}
            lines.append(target_line)
        words_list = target_line["words"]
        assert isinstance(words_list, list)
        words_list.append(word)

    normalized_lines: list[list[PdfWordBox]] = []
    for line in sorted(lines, key=lambda item: float(item["y_mid"])):
        words_list = line["words"]
        assert isinstance(words_list, list)
        normalized_lines.append(sorted(words_list, key=lambda item: item.x0))
    return normalized_lines


def _separator_between_words(previous: PdfWordBox, current: PdfWordBox) -> str:
    gap = current.x0 - previous.x1
    if previous.text in PIPE_TOKENS or current.text in PIPE_TOKENS:
        return " "
    if gap >= PDF_TABLE_GAP_THRESHOLD and len(previous.text) <= 24 and len(current.text) <= 48:
        return " | "
    if gap >= PDF_WORD_GAP_THRESHOLD:
        return " "
    return ""


def _normalize_pdftotext_page(words: list[PdfWordBox], page_index: int) -> PdfNormalizedPage:
    if not words:
        return PdfNormalizedPage(text="", char_boxes=[], provider="pdftotext")

    parts: list[str] = []
    char_boxes: list[PdfRectBox | None] = []
    lines = _group_words_into_lines(words)

    for line_index, line_words in enumerate(lines):
        for word_index, word in enumerate(line_words):
            if word_index > 0:
                _append_literal(parts, char_boxes, _separator_between_words(line_words[word_index - 1], word))
            _append_positioned_text(parts, char_boxes, word.text, _distribute_word_boxes(word, page_index))
        if line_index < len(lines) - 1:
            _append_literal(parts, char_boxes, "\n")

    return PdfNormalizedPage(text="".join(parts), char_boxes=char_boxes, provider="pdftotext")


def _normalize_pymupdf_page(page: "fitz.Page", page_index: int) -> PdfNormalizedPage:
    raw = page.get_text("rawdict")
    parts: list[str] = []
    char_boxes: list[PdfRectBox | None] = []

    for block in raw.get("blocks", []):
        for line in block.get("lines", []):
            line_has_chars = False
            for span in line.get("spans", []):
                for char in span.get("chars", []):
                    value = char.get("c", "")
                    if not value:
                        continue
                    bbox = char.get("bbox")
                    parts.append(value)
                    if isinstance(bbox, (list, tuple)) and len(bbox) == 4:
                        char_boxes.append(
                            PdfRectBox(
                                page_index,
                                float(bbox[0]),
                                float(bbox[1]),
                                float(bbox[2]),
                                float(bbox[3]),
                            )
                        )
                    else:
                        char_boxes.append(None)
                    line_has_chars = True
            if line_has_chars:
                _append_literal(parts, char_boxes, "\n")

    if parts and parts[-1] == "\n":
        parts.pop()
        char_boxes.pop()

    return PdfNormalizedPage(text="".join(parts), char_boxes=char_boxes, provider="pymupdf")


def _normalize_simple_ocr_text(text: str) -> PdfNormalizedPage:
    return PdfNormalizedPage(text=text.strip(), char_boxes=[None] * len(text.strip()), provider="vision_ocr")


def _bbox_from_ocr_payload(
    bbox: dict[str, object],
    page_index: int,
    page_width: float,
    page_height: float,
) -> PdfRectBox | None:
    try:
        x0 = float(bbox["x0"]) * page_width
        x1 = float(bbox["x1"]) * page_width
        y0 = (1.0 - float(bbox["y1"])) * page_height
        y1 = (1.0 - float(bbox["y0"])) * page_height
    except (KeyError, TypeError, ValueError):
        return None
    return PdfRectBox(page_index, x0, y0, x1, y1)


def _normalize_vision_ocr_page(payload: object, page_index: int, page_width: float, page_height: float) -> PdfNormalizedPage:
    if isinstance(payload, str):
        return _normalize_simple_ocr_text(payload)

    if not isinstance(payload, dict):
        return PdfNormalizedPage(text="", char_boxes=[], provider="vision_ocr")

    raw_lines = payload.get("lines", [])
    if not isinstance(raw_lines, list):
        return PdfNormalizedPage(text="", char_boxes=[], provider="vision_ocr")

    def _line_sort_key(line: dict[str, object]) -> tuple[float, float]:
        bbox = line.get("bbox", {})
        if not isinstance(bbox, dict):
            return (1.0, 0.0)
        try:
            top = 1.0 - float(bbox["y1"])
            left = float(bbox["x0"])
        except (KeyError, TypeError, ValueError):
            return (1.0, 0.0)
        return (top, left)

    parts: list[str] = []
    char_boxes: list[PdfRectBox | None] = []
    ordered_lines = sorted((line for line in raw_lines if isinstance(line, dict)), key=_line_sort_key)
    for line_index, line in enumerate(ordered_lines):
        raw_chars = line.get("chars", [])
        if not isinstance(raw_chars, list):
            raw_chars = []
        line_text_parts: list[str] = []
        line_char_boxes: list[PdfRectBox | None] = []
        for raw_char in raw_chars:
            if not isinstance(raw_char, dict):
                continue
            char_text = str(raw_char.get("text", ""))
            if not char_text:
                continue
            for fragment in char_text:
                bbox = raw_char.get("bbox")
                rect = _bbox_from_ocr_payload(bbox, page_index, page_width, page_height) if isinstance(bbox, dict) else None
                line_text_parts.append(fragment)
                line_char_boxes.append(None if fragment.isspace() else rect)

        line_text = "".join(line_text_parts).strip()
        if not line_text:
            continue

        leading_trim = len("".join(line_text_parts)) - len("".join(line_text_parts).lstrip())
        trailing_trim = len("".join(line_text_parts)) - len("".join(line_text_parts).rstrip())
        if leading_trim:
            line_char_boxes = line_char_boxes[leading_trim:]
        if trailing_trim:
            line_char_boxes = line_char_boxes[: len(line_char_boxes) - trailing_trim]

        _append_positioned_text(parts, char_boxes, line_text, line_char_boxes)
        if line_index < len(ordered_lines) - 1:
            _append_literal(parts, char_boxes, "\n")

    return PdfNormalizedPage(text="".join(parts), char_boxes=char_boxes, provider="vision_ocr")


def extract_pdf_with_metadata(path: str, skip_ocr: bool = False) -> PdfExtractionResult:
    try:
        import fitz
    except ImportError as exc:
        raise UnsupportedFileTypeError("PyMuPDF not installed; cannot extract PDF") from exc

    doc = fitz.open(path)
    try:
        pdftotext_pages = _load_pdftotext_word_pages(path) or []
        normalized_pages: list[PdfNormalizedPage] = []
        page_flags: list[set[str]] = []
        ocr_page_numbers: list[int] = []
        text_layer_char_count = 0
        ocr_char_count = 0

        for page_index, page in enumerate(doc):
            raw_text = page.get_text()
            flags = _collect_pdf_quality_flags(page, raw_text)
            page_flags.append(flags)

            if page_index < len(pdftotext_pages) and pdftotext_pages[page_index]:
                normalized_page = _normalize_pdftotext_page(pdftotext_pages[page_index], page_index)
                if not normalized_page.text.strip() and raw_text.strip():
                    normalized_page = _normalize_pymupdf_page(page, page_index)
            elif raw_text.strip():
                normalized_page = _normalize_pymupdf_page(page, page_index)
            else:
                normalized_page = PdfNormalizedPage(text="", char_boxes=[], provider="pymupdf")

            normalized_pages.append(normalized_page)
            text_layer_char_count += len(normalized_page.text)

            if "no_text_layer" in flags or "low_text_density" in flags:
                ocr_page_numbers.append(page_index)

        if ocr_page_numbers and not skip_ocr:
            try:
                ocr_pages = _extract_pdf_with_ocr(path, ocr_page_numbers)
            except (ImportError, OSError, RuntimeError, UnsupportedFileTypeError):
                ocr_pages = []

            for page_index, ocr_payload in zip(ocr_page_numbers, ocr_pages):
                page = doc.load_page(page_index)
                normalized_page = _normalize_vision_ocr_page(
                    ocr_payload,
                    page_index,
                    float(page.rect.width),
                    float(page.rect.height),
                )
                if not normalized_page.text.strip():
                    continue
                normalized_pages[page_index] = normalized_page
                ocr_char_count += len(normalized_page.text)

        page_texts = [page.text for page in normalized_pages]
        page_providers = [page.provider for page in normalized_pages]
        char_boxes: list[PdfRectBox | None] = []
        for page_index, page in enumerate(normalized_pages):
            char_boxes.extend(page.char_boxes)
            if page_index < len(normalized_pages) - 1:
                char_boxes.extend([None] * len(PDF_PAGE_SEPARATOR))

        pdf_provider = page_providers[0] if page_providers and len(set(page_providers)) == 1 else "hybrid"
        quality_flags = sorted({flag for flags in page_flags for flag in flags})
        used_ocr = any(provider == "vision_ocr" for provider in page_providers)
        return PdfExtractionResult(
            text=PDF_PAGE_SEPARATOR.join(page_texts),
            page_count=doc.page_count,
            used_ocr=used_ocr,
            text_layer_char_count=text_layer_char_count,
            ocr_char_count=ocr_char_count,
            quality_flags=quality_flags,
            pdf_provider=pdf_provider,
            page_providers=page_providers,
            char_boxes=char_boxes,
        )
    finally:
        doc.close()


def _extract_pdf(path: str) -> str:
    return extract_pdf_with_metadata(path).text


def _extract_docx(path: str) -> str:
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise UnsupportedFileTypeError("python-docx not installed; cannot extract DOCX") from exc

    doc = Document(path)
    parts: list[str] = []

    def _collect_table(tbl: "Table") -> None:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)
                for nested in cell.tables:
                    _collect_table(nested)

    for child in doc.element.body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local == "p":
            text = Paragraph(child, doc).text.strip()
            if text:
                parts.append(text)
        elif local == "tbl":
            _collect_table(Table(child, doc))

    return "\n".join(parts)


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
}


def extract_text(path: str, suffix: str | None = None) -> str:
    ext = (suffix or Path(path).suffix).lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'; supported: {', '.join(_EXTRACTORS)}"
        )
    return extractor(path)


def _boxes_on_same_line(left: PdfRectBox, right: PdfRectBox) -> bool:
    if left.page_index != right.page_index:
        return False
    overlap = max(0.0, min(left.y1, right.y1) - max(left.y0, right.y0))
    min_height = min(left.y1 - left.y0, right.y1 - right.y0)
    return min_height > 0 and overlap / min_height >= 0.5


def _merge_rect_boxes(boxes: list[PdfRectBox]) -> list[PdfRectBox]:
    if not boxes:
        return []
    ordered = sorted(boxes, key=lambda box: (box.page_index, box.y0, box.x0, box.x1))
    merged: list[PdfRectBox] = [ordered[0]]
    for box in ordered[1:]:
        previous = merged[-1]
        if _boxes_on_same_line(previous, box) and box.x0 - previous.x1 <= PDF_RECT_GAP_THRESHOLD:
            merged[-1] = PdfRectBox(
                previous.page_index,
                min(previous.x0, box.x0),
                min(previous.y0, box.y0),
                max(previous.x1, box.x1),
                max(previous.y1, box.y1),
            )
            continue
        merged.append(box)
    return merged


def _find_value_rects(page: "fitz.Page", value: str) -> list["fitz.Rect"]:
    import fitz

    rects = page.search_for(value)
    if rects:
        return rects
    return []


def redact_pdf_bytes(
    path: str,
    extraction: PdfExtractionResult | list[str],
    spans: list[dict[str, object]] | None = None,
) -> bytes:
    try:
        import fitz
    except ImportError as exc:
        raise UnsupportedFileTypeError("PyMuPDF not installed; cannot redact PDF") from exc

    if isinstance(extraction, PdfExtractionResult):
        active_spans = spans or []
        page_rects: dict[int, list[PdfRectBox]] = {}
        # Spans with no char_box coordinates (e.g. OCR pages extracted with skip_ocr=True)
        fallback_values: set[str] = set()
        for span in active_spans:
            rects = _merge_rect_boxes(
                extraction.rect_boxes_for_span(int(span["start"]), int(span["end"]))
            )
            if rects:
                for rect in rects:
                    page_rects.setdefault(rect.page_index, []).append(rect)
            else:
                span_text = str(span.get("text", "")).strip()
                if span_text:
                    fallback_values.add(span_text)

        if not page_rects and not fallback_values:
            return Path(path).read_bytes()

        doc = fitz.open(path)
        try:
            for page_index, rects in page_rects.items():
                page = doc.load_page(page_index)
                for rect in rects:
                    page.add_redact_annot(
                        fitz.Rect(rect.x0, rect.y0, rect.x1, rect.y1),
                        fill=(0, 0, 0),
                    )
                page.apply_redactions()
            # Fallback: use text search for spans that had no coordinate data
            if fallback_values:
                for page in doc:
                    for value in fallback_values:
                        for rect in page.search_for(value):
                            page.add_redact_annot(rect, fill=(0, 0, 0))
                    page.apply_redactions()
            return doc.tobytes(garbage=4, deflate=True)
        finally:
            doc.close()

    secret_values = extraction
    doc = fitz.open(path)
    try:
        for page in doc:
            for value in secret_values:
                if not value:
                    continue
                for rect in _find_value_rects(page, value):
                    page.add_redact_annot(rect, fill=(0, 0, 0))
            page.apply_redactions()
        return doc.tobytes(garbage=4, deflate=True)
    finally:
        doc.close()


def redact_docx_bytes(path: str, secret_values: list[str]) -> bytes:
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise UnsupportedFileTypeError("python-docx not installed; cannot redact DOCX") from exc

    doc = Document(path)
    values = [value for value in secret_values if value]

    def _redact_para(para: "Paragraph") -> None:
        for value in values:
            for run in para.runs:
                if value in run.text:
                    run.text = run.text.replace(value, "<SECRET>")
            if value in para.text:
                replaced = para.text.replace(value, "<SECRET>")
                runs = para.runs
                if runs:
                    runs[0].text = replaced
                    for run in runs[1:]:
                        run.text = ""

    def _redact_table(tbl: "Table") -> None:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _redact_para(para)
                for nested in cell.tables:
                    _redact_table(nested)

    for child in doc.element.body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local == "p":
            _redact_para(Paragraph(child, doc))
        elif local == "tbl":
            _redact_table(Table(child, doc))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
