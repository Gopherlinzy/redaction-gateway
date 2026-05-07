"""
Tests for POST /redact-file endpoint (PDF + DOCX extraction + redaction).
Uses in-memory synthetic files — no real documents needed.
"""
import io
from pathlib import Path
from fastapi.testclient import TestClient
import pytest

import app as app_module
from detectors.file_extractor import extract_pdf_with_metadata

client = TestClient(app_module.app)

SECRET_LINE = "OPENAI_API_KEY=sk-testsecretvalue1234567890abcdef"
CLEAN_LINE  = "Hello world, this is a normal document."
PERSON_LINE = "姓名 | 孙禄毅"
FIXTURE_DIR = Path(__file__).resolve().parent / "fixtures" / "pdf"


def _fixture_bytes(name: str) -> bytes:
    return (FIXTURE_DIR / name).read_bytes()


def _fixture_path(name: str) -> str:
    return str(FIXTURE_DIR / name)


def _make_pdf_bytes(text: str) -> bytes:
    """Build a minimal single-page PDF containing `text`."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    out = doc.tobytes()
    doc.close()
    return out


def _make_multipage_pdf_bytes(*pages_text: str) -> bytes:
    import fitz

    doc = fitz.open()
    for text in pages_text:
        page = doc.new_page()
        page.insert_text((72, 72), text)
    out = doc.tobytes()
    doc.close()
    return out


def _make_scanned_pdf_bytes(text: str) -> bytes:
    """
    Build an image-only PDF by rasterizing a text PDF page and embedding the
    rendered page as an image into a fresh PDF.
    """
    import fitz

    source = fitz.open()
    source_page = source.new_page()
    source_page.insert_text((72, 72), text, fontsize=14)
    pix = source_page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    png_bytes = pix.tobytes("png")

    scanned = fitz.open()
    page = scanned.new_page(width=pix.width, height=pix.height)
    page.insert_image(fitz.Rect(0, 0, pix.width, pix.height), stream=png_bytes)

    out = scanned.tobytes()
    source.close()
    scanned.close()
    return out


def _make_low_density_pdf_bytes(secret_text: str) -> bytes:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Invoice ID")
    page.insert_text((72, 140), secret_text, fontsize=1)
    out = doc.tobytes()
    doc.close()
    return out


def _make_docx_bytes(text: str) -> bytes:
    """Build a minimal DOCX with a single paragraph containing `text`."""
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _post_file(content: bytes, filename: str, detection_mode: str = "balanced") -> dict:
    resp = client.post(
        "/redact-file",
        files={"file": (filename, content, "application/octet-stream")},
        data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
              "detection_mode": detection_mode},
    )
    assert resp.status_code == 200, resp.text
    return resp.json()


def _install_fake_person_opf(monkeypatch) -> None:
    def _fake_detect_with_runtime(text: str):
        if "孙禄毅" not in text:
            return []
        start = text.index("孙禄毅")
        end = start + len("孙禄毅")
        return [{
            "label": "private_person",
            "start": start,
            "end": end,
            "text": "孙禄毅",
            "placeholder": "<PRIVATE_PERSON>",
            "source": "opf",
        }]

    monkeypatch.setattr(app_module, "detect_with_runtime", _fake_detect_with_runtime)


def _install_fake_ocr(monkeypatch, text: str) -> None:
    from detectors import file_extractor

    def _fake_extract_pdf_with_ocr(path: str, page_numbers: list[int] | None = None) -> list[str]:
        import fitz

        doc = fitz.open(path)
        try:
            indexes = page_numbers if page_numbers is not None else list(range(doc.page_count))
            return [text for _ in indexes]
        finally:
            doc.close()

    monkeypatch.setattr(
        file_extractor,
        "_extract_pdf_with_ocr",
        _fake_extract_pdf_with_ocr,
        raising=False,
    )


def _disable_ocr(monkeypatch) -> None:
    from detectors import file_extractor

    def _failing_extract_pdf_with_ocr(path: str, page_numbers: list[int] | None = None) -> list[str]:
        raise RuntimeError("OCR disabled for test")

    monkeypatch.setattr(
        file_extractor,
        "_extract_pdf_with_ocr",
        _failing_extract_pdf_with_ocr,
        raising=False,
    )


# ── PDF tests ────────────────────────────────────────────────────────────────

class TestRedactFilePDF:
    def test_pdf_with_secret_is_detected(self):
        body = _post_file(_make_pdf_bytes(SECRET_LINE), "test.pdf")
        assert body["file_type"] == "pdf"
        assert body["summary"]["secret_count"] >= 1
        assert "<SECRET>" in body["redacted_text"]
        assert "sk-testsecretvalue1234567890abcdef" not in body["redacted_text"]

    def test_pdf_without_secret_passes_clean(self):
        body = _post_file(_make_pdf_bytes(CLEAN_LINE), "clean.pdf")
        assert body["summary"]["secret_count"] == 0
        assert "<SECRET>" not in body["redacted_text"]

    def test_pdf_returns_metadata_and_text_layer_extraction_fields(self):
        body = _post_file(_make_pdf_bytes(CLEAN_LINE), "meta.pdf")
        assert body["filename"] == "meta.pdf"
        assert body["file_type"] == "pdf"
        assert isinstance(body["char_count"], int)
        assert body["char_count"] > 0
        assert CLEAN_LINE in body["source_text"]
        assert body["extraction_method"] == "text_layer"
        assert isinstance(body["extraction_flags"], list)
        assert "no_text_layer" not in body["extraction_flags"]

    def test_pdf_redacted_text_preserves_clean_content(self):
        body = _post_file(_make_pdf_bytes(CLEAN_LINE), "preserve.pdf")
        assert CLEAN_LINE in body["redacted_text"]

    def test_scanned_pdf_uses_ocr_fallback_when_available(self, monkeypatch):
        _install_fake_ocr(monkeypatch, SECRET_LINE)
        body = _post_file(_make_scanned_pdf_bytes(SECRET_LINE), "scan.pdf")
        assert body["file_type"] == "pdf"
        assert body["extraction_method"] == "ocr"
        assert "no_text_layer" in body["extraction_flags"]
        assert body["summary"]["secret_count"] >= 1
        assert "<SECRET>" in body["redacted_text"]
        assert "sk-testsecretvalue1234567890abcdef" not in body["redacted_text"]

    def test_low_density_pdf_uses_ocr_fallback_when_available(self, monkeypatch):
        _install_fake_ocr(monkeypatch, SECRET_LINE)
        body = _post_file(_make_low_density_pdf_bytes(SECRET_LINE), "sparse.pdf")
        assert body["extraction_method"] == "ocr"
        assert "low_text_density" in body["extraction_flags"]
        assert body["summary"]["secret_count"] >= 1

    def test_clean_scanned_pdf_stays_clean_with_ocr(self, monkeypatch):
        _install_fake_ocr(monkeypatch, CLEAN_LINE)
        body = _post_file(_make_scanned_pdf_bytes(CLEAN_LINE), "clean-scan.pdf")
        assert body["extraction_method"] == "ocr"
        assert body["summary"]["secret_count"] == 0
        assert CLEAN_LINE in body["redacted_text"]

    def test_multipage_pdf_preserves_page_order_in_extracted_text(self):
        page_one = "Page one line"
        page_two = SECRET_LINE
        body = _post_file(_make_multipage_pdf_bytes(page_one, page_two), "multipage.pdf")
        assert body["source_text"].index(page_one) < body["source_text"].index(page_two)
        assert "\n\f\n" in body["source_text"]
        assert body["summary"]["secret_count"] >= 1

    def test_short_text_backed_pdf_still_runs_opf_for_person_detection(self, monkeypatch):
        _install_fake_person_opf(monkeypatch)
        body = _post_file(_fixture_bytes("text_backed_cjk_table.pdf"), "text_backed_cjk_table.pdf")
        assert any(span["label"] == "private_person" for span in body["spans"])
        assert "<PRIVATE_PERSON>" in body["redacted_text"]

    def test_text_backed_pdf_uses_pdftotext_provider_and_normalizes_table_rows(self):
        extraction = extract_pdf_with_metadata(_fixture_path("text_backed_cjk_table.pdf"))
        assert extraction.pdf_provider == "pdftotext"
        assert extraction.page_providers == ["pdftotext"]
        assert PERSON_LINE in extraction.text

    def test_mixed_layout_pdf_inserts_structured_separators_for_table_cells(self):
        extraction = extract_pdf_with_metadata(_fixture_path("mixed_layout_cjk.pdf"))
        assert "姓名 | 孙禄毅" in extraction.text
        assert "部门 | 风控部" in extraction.text

    def test_scanned_pdf_uses_vision_ocr_provider_and_restores_person_line(self):
        extraction = extract_pdf_with_metadata(_fixture_path("scanned_cjk_table.pdf"))
        if extraction.used_ocr is False:
            pytest.skip("Vision OCR runtime unavailable in current execution environment")
        assert extraction.used_ocr is True
        assert "vision_ocr" in extraction.page_providers
        assert PERSON_LINE in extraction.text

    def test_text_backed_pdf_person_field_fallback_detects_without_opf(self, monkeypatch):
        monkeypatch.setattr(app_module, "detect_with_runtime", lambda _text: [])
        body = _post_file(_fixture_bytes("text_backed_cjk_table.pdf"), "text_backed_cjk_table.pdf")
        assert any(
            span["label"] == "private_person" and span["source"] == "regex"
            for span in body["spans"]
        )


# ── DOCX tests ───────────────────────────────────────────────────────────────

class TestRedactFileDOCX:
    def test_docx_with_secret_is_detected(self):
        body = _post_file(_make_docx_bytes(SECRET_LINE), "test.docx")
        assert body["file_type"] == "docx"
        assert body["summary"]["secret_count"] >= 1
        assert "<SECRET>" in body["redacted_text"]
        assert "sk-testsecretvalue1234567890abcdef" not in body["redacted_text"]

    def test_docx_without_secret_passes_clean(self):
        body = _post_file(_make_docx_bytes(CLEAN_LINE), "clean.docx")
        assert body["summary"]["secret_count"] == 0

    def test_docx_returns_metadata(self):
        body = _post_file(_make_docx_bytes(CLEAN_LINE), "meta.docx")
        assert body["filename"] == "meta.docx"
        assert body["file_type"] == "docx"
        assert body["char_count"] > 0
        assert CLEAN_LINE in body["source_text"]

    def test_short_docx_still_runs_opf_for_person_detection(self, monkeypatch):
        _install_fake_person_opf(monkeypatch)
        body = _post_file(_make_docx_bytes(PERSON_LINE), "person.docx")
        assert any(span["label"] == "private_person" for span in body["spans"])
        assert "<PRIVATE_PERSON>" in body["redacted_text"]


# ── Error handling tests ─────────────────────────────────────────────────────

class TestRedactFileErrors:
    def test_unsupported_extension_returns_422(self):
        resp = client.post(
            "/redact-file",
            files={"file": ("secret.txt", b"OPENAI_API_KEY=sk-test", "text/plain")},
            data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
                  "detection_mode": "balanced"},
        )
        assert resp.status_code == 422
        assert "Unsupported" in resp.json()["error"]

    def test_empty_pdf_returns_422(self):
        import fitz
        doc = fitz.open()
        doc.new_page()  # blank page, no text
        pdf_bytes = doc.tobytes()
        resp = client.post(
            "/redact-file",
            files={"file": ("empty.pdf", pdf_bytes, "application/pdf")},
            data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
                  "detection_mode": "balanced"},
        )
        assert resp.status_code == 422
        assert "No text" in resp.json()["error"]

    def test_scanned_pdf_without_ocr_and_without_text_returns_422(self, monkeypatch):
        _disable_ocr(monkeypatch)
        resp = client.post(
            "/redact-file",
            files={"file": ("scan.pdf", _make_scanned_pdf_bytes(SECRET_LINE), "application/pdf")},
            data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
                  "detection_mode": "balanced"},
        )
        assert resp.status_code == 422
        assert "No text" in resp.json()["error"]


# ── /redact-file/download tests ───────────────────────────────────────────────

def _post_download(content: bytes, filename: str) -> "Response":
    return client.post(
        "/redact-file/download",
        files={"file": (filename, content, "application/octet-stream")},
        data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
              "detection_mode": "balanced"},
    )


class TestRedactFileDownloadPDF:
    def test_returns_pdf_content_type(self):
        resp = _post_download(_make_pdf_bytes(SECRET_LINE), "report.pdf")
        assert resp.status_code == 200
        assert "application/pdf" in resp.headers["content-type"]

    def test_filename_has_redacted_suffix(self):
        resp = _post_download(_make_pdf_bytes(SECRET_LINE), "report.pdf")
        assert "report_redacted.pdf" in resp.headers["content-disposition"]

    def test_secret_not_present_in_returned_pdf_text(self):
        import fitz
        resp = _post_download(_make_pdf_bytes(SECRET_LINE), "secret.pdf")
        assert resp.status_code == 200
        doc = fitz.open(stream=resp.content, filetype="pdf")
        full_text = "".join(page.get_text() for page in doc)
        assert "sk-testsecretvalue1234567890abcdef" not in full_text

    def test_clean_pdf_passes_through_unchanged_content(self):
        import fitz
        resp = _post_download(_make_pdf_bytes(CLEAN_LINE), "clean.pdf")
        assert resp.status_code == 200
        doc = fitz.open(stream=resp.content, filetype="pdf")
        full_text = "".join(page.get_text() for page in doc)
        assert CLEAN_LINE in full_text

    def test_scanned_pdf_download_still_succeeds_with_ocr_detection(self, monkeypatch):
        _install_fake_ocr(monkeypatch, SECRET_LINE)
        resp = _post_download(_make_scanned_pdf_bytes(SECRET_LINE), "scan.pdf")
        assert resp.status_code == 200
        assert "application/pdf" in resp.headers["content-type"]

    def test_text_backed_pdf_download_redacts_person_using_provider_coords(self, monkeypatch):
        import fitz

        _install_fake_person_opf(monkeypatch)
        resp = _post_download(_fixture_bytes("text_backed_cjk_table.pdf"), "text_backed_cjk_table.pdf")
        assert resp.status_code == 200
        doc = fitz.open(stream=resp.content, filetype="pdf")
        full_text = "".join(page.get_text() for page in doc)
        assert "孙禄毅" not in full_text
        assert "姓名" in full_text


class TestRedactFileDownloadDOCX:
    def test_returns_docx_content_type(self):
        resp = _post_download(_make_docx_bytes(SECRET_LINE), "report.docx")
        assert resp.status_code == 200
        assert "wordprocessingml" in resp.headers["content-type"]

    def test_filename_has_redacted_suffix(self):
        resp = _post_download(_make_docx_bytes(SECRET_LINE), "report.docx")
        assert "report_redacted.docx" in resp.headers["content-disposition"]

    def test_secret_replaced_in_returned_docx(self):
        from docx import Document
        resp = _post_download(_make_docx_bytes(SECRET_LINE), "secret.docx")
        assert resp.status_code == 200
        doc = Document(io.BytesIO(resp.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "sk-testsecretvalue1234567890abcdef" not in full_text
        assert "<SECRET>" in full_text

    def test_clean_docx_passes_through(self):
        from docx import Document
        resp = _post_download(_make_docx_bytes(CLEAN_LINE), "clean.docx")
        assert resp.status_code == 200
        doc = Document(io.BytesIO(resp.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert CLEAN_LINE in full_text

    def test_download_unsupported_type_returns_422(self):
        resp = _post_download(b"secret content", "data.txt")
        assert resp.status_code == 422


# ── active_spans precision tests ─────────────────────────────────────────────

SECRET_A = "OPENAI_API_KEY=sk-aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa"
SECRET_B = "GITHUB_TOKEN=ghp_bbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbb"


def _make_docx_two_secrets() -> bytes:
    """DOCX with two different secrets in separate paragraphs."""
    from docx import Document
    doc = Document()
    doc.add_paragraph(SECRET_A)
    doc.add_paragraph(SECRET_B)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf_two_secrets() -> bytes:
    """Text-backed PDF with two secrets on separate lines."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), SECRET_A)
    page.insert_text((72, 120), SECRET_B)
    out = doc.tobytes()
    doc.close()
    return out


def _post_scan(content: bytes, filename: str) -> dict:
    resp = client.post(
        "/redact-file",
        files={"file": (filename, content, "application/octet-stream")},
        data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
              "detection_mode": "balanced"},
    )
    assert resp.status_code == 200, resp.text
    return resp.json()


def _post_download_with_active_spans(content: bytes, filename: str, active_spans: list) -> "Response":
    import json
    return client.post(
        "/redact-file/download",
        files={"file": (filename, content, "application/octet-stream")},
        data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
              "detection_mode": "balanced",
              "active_spans": json.dumps(active_spans)},
    )


def _post_download_with_active_categories(content: bytes, filename: str, active_categories: list) -> "Response":
    import json
    return client.post(
        "/redact-file/download",
        files={"file": (filename, content, "application/octet-stream")},
        data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
              "detection_mode": "balanced",
              "active_categories": json.dumps(active_categories)},
    )


class TestActiveSpansDocx:
    def test_disable_one_of_two_same_label_spans_restores_only_that_one(self):
        """Toggling off span A must not affect span B."""
        from docx import Document

        content = _make_docx_two_secrets()
        scan = _post_scan(content, "two.docx")
        spans = scan["spans"]
        source_text = scan["source_text"]

        # Identify spans for SECRET_A and SECRET_B by offset
        secret_a_val = "sk-aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa"
        secret_b_val = "ghp_bbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbb"

        spans_a = [s for s in spans if source_text[s["start"]:s["end"]] == secret_a_val]
        spans_b = [s for s in spans if source_text[s["start"]:s["end"]] == secret_b_val]
        assert spans_a, "span for SECRET_A not detected"
        assert spans_b, "span for SECRET_B not detected"

        # Keep only SECRET_B active (disable SECRET_A)
        active = [{"start": s["start"], "end": s["end"], "label": s["label"]} for s in spans_b]
        resp = _post_download_with_active_spans(content, "two.docx", active)
        assert resp.status_code == 200

        doc = Document(io.BytesIO(resp.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert secret_a_val in full_text, "disabled span should be restored"
        assert secret_b_val not in full_text, "enabled span should remain redacted"

    def test_backward_compat_active_categories_still_works(self):
        """Passing only active_categories (no active_spans) still redacts by category."""
        from docx import Document

        content = _make_docx_two_secrets()
        resp = _post_download_with_active_categories(content, "cats.docx", ["secret"])
        assert resp.status_code == 200

        doc = Document(io.BytesIO(resp.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "<SECRET>" in full_text


class TestActiveSpansPdf:
    def test_disable_one_of_two_spans_restores_only_that_one_in_pdf(self):
        """Text-backed PDF: toggling off span A keeps span B redacted."""
        import fitz

        content = _make_pdf_two_secrets()
        scan = _post_scan(content, "two.pdf")
        spans = scan["spans"]
        source_text = scan["source_text"]

        secret_a_val = "sk-aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa"
        secret_b_val = "ghp_bbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbb"

        spans_a = [s for s in spans if source_text[s["start"]:s["end"]] == secret_a_val]
        spans_b = [s for s in spans if source_text[s["start"]:s["end"]] == secret_b_val]
        assert spans_a, "span for SECRET_A not detected in PDF"
        assert spans_b, "span for SECRET_B not detected in PDF"

        # Keep only SECRET_B active
        active = [{"start": s["start"], "end": s["end"], "label": s["label"]} for s in spans_b]
        resp = _post_download_with_active_spans(content, "two.pdf", active)
        assert resp.status_code == 200

        doc = fitz.open(stream=resp.content, filetype="pdf")
        full_text = "".join(page.get_text() for page in doc)

        assert secret_a_val in full_text, "disabled PDF span should be visible"
        assert secret_b_val not in full_text, "enabled PDF span should remain redacted"

    def test_empty_active_spans_list_passes_through_all_text(self):
        """active_spans=[] means nothing is active — file should pass through unredacted."""
        import fitz

        content = _make_pdf_two_secrets()
        resp = _post_download_with_active_spans(content, "none.pdf", [])
        assert resp.status_code == 200

        doc = fitz.open(stream=resp.content, filetype="pdf")
        full_text = "".join(page.get_text() for page in doc)
        assert "sk-aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa" in full_text
        assert "ghp_bbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbb" in full_text
