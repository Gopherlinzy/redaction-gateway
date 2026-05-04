"""
Tests for POST /redact-file endpoint (PDF + DOCX extraction + redaction).
Uses in-memory synthetic files — no real documents needed.
"""
import io
import pytest
from fastapi.testclient import TestClient

from app import app

client = TestClient(app)

SECRET_LINE = "OPENAI_API_KEY=sk-testsecretvalue1234567890abcdef"
CLEAN_LINE  = "Hello world, this is a normal document."


def _make_pdf_bytes(text: str) -> bytes:
    """Build a minimal single-page PDF containing `text`."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    return doc.tobytes()


def _make_docx_bytes(text: str) -> bytes:
    """Build a minimal DOCX with a single paragraph containing `text`."""
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _post_file(content: bytes, filename: str, detection_mode: str = "balanced") -> dict:
    resp = client.post(
        "/redact-file",
        files={"file": (filename, content, "application/octet-stream")},
        data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
              "detection_mode": detection_mode},
    )
    assert resp.status_code == 200, resp.text
    return resp.json()


# ── PDF tests ────────────────────────────────────────────────────────────────

class TestRedactFilePDF:
    def test_pdf_with_secret_is_detected(self):
        body = _post_file(_make_pdf_bytes(SECRET_LINE), "test.pdf")
        assert body["file_type"] == "pdf"
        assert body["summary"]["secret_count"] >= 1
        assert "<SECRET>" in body["redacted_text"]
        assert "sk-testsecretvalue1234567890abcdef" not in body["redacted_text"]

    def test_pdf_without_secret_passes_clean(self):
        body = _post_file(_make_pdf_bytes(CLEAN_LINE), "clean.pdf")
        assert body["summary"]["secret_count"] == 0
        assert "<SECRET>" not in body["redacted_text"]

    def test_pdf_returns_metadata(self):
        body = _post_file(_make_pdf_bytes(CLEAN_LINE), "meta.pdf")
        assert body["filename"] == "meta.pdf"
        assert body["file_type"] == "pdf"
        assert isinstance(body["char_count"], int)
        assert body["char_count"] > 0
        assert CLEAN_LINE in body["source_text"]

    def test_pdf_redacted_text_preserves_clean_content(self):
        body = _post_file(_make_pdf_bytes(CLEAN_LINE), "preserve.pdf")
        assert CLEAN_LINE in body["redacted_text"]


# ── DOCX tests ───────────────────────────────────────────────────────────────

class TestRedactFileDOCX:
    def test_docx_with_secret_is_detected(self):
        body = _post_file(_make_docx_bytes(SECRET_LINE), "test.docx")
        assert body["file_type"] == "docx"
        assert body["summary"]["secret_count"] >= 1
        assert "<SECRET>" in body["redacted_text"]
        assert "sk-testsecretvalue1234567890abcdef" not in body["redacted_text"]

    def test_docx_without_secret_passes_clean(self):
        body = _post_file(_make_docx_bytes(CLEAN_LINE), "clean.docx")
        assert body["summary"]["secret_count"] == 0

    def test_docx_returns_metadata(self):
        body = _post_file(_make_docx_bytes(CLEAN_LINE), "meta.docx")
        assert body["filename"] == "meta.docx"
        assert body["file_type"] == "docx"
        assert body["char_count"] > 0
        assert CLEAN_LINE in body["source_text"]


# ── Error handling tests ─────────────────────────────────────────────────────

class TestRedactFileErrors:
    def test_unsupported_extension_returns_422(self):
        resp = client.post(
            "/redact-file",
            files={"file": ("secret.txt", b"OPENAI_API_KEY=sk-test", "text/plain")},
            data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
                  "detection_mode": "balanced"},
        )
        assert resp.status_code == 422
        assert "Unsupported" in resp.json()["error"]

    def test_empty_pdf_returns_422(self):
        import fitz
        doc = fitz.open()
        doc.new_page()  # blank page, no text
        pdf_bytes = doc.tobytes()
        resp = client.post(
            "/redact-file",
            files={"file": ("empty.pdf", pdf_bytes, "application/pdf")},
            data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
                  "detection_mode": "balanced"},
        )
        assert resp.status_code == 422
        assert "No text" in resp.json()["error"]


# ── /redact-file/download tests ───────────────────────────────────────────────

def _post_download(content: bytes, filename: str) -> "Response":
    return client.post(
        "/redact-file/download",
        files={"file": (filename, content, "application/octet-stream")},
        data={"source": "manual_ui", "target": "ai_model", "mode": "warn",
              "detection_mode": "balanced"},
    )


class TestRedactFileDownloadPDF:
    def test_returns_pdf_content_type(self):
        resp = _post_download(_make_pdf_bytes(SECRET_LINE), "report.pdf")
        assert resp.status_code == 200
        assert "application/pdf" in resp.headers["content-type"]

    def test_filename_has_redacted_suffix(self):
        resp = _post_download(_make_pdf_bytes(SECRET_LINE), "report.pdf")
        assert "report_redacted.pdf" in resp.headers["content-disposition"]

    def test_secret_not_present_in_returned_pdf_text(self):
        import fitz
        resp = _post_download(_make_pdf_bytes(SECRET_LINE), "secret.pdf")
        assert resp.status_code == 200
        doc = fitz.open(stream=resp.content, filetype="pdf")
        full_text = "".join(page.get_text() for page in doc)
        assert "sk-testsecretvalue1234567890abcdef" not in full_text

    def test_clean_pdf_passes_through_unchanged_content(self):
        import fitz
        resp = _post_download(_make_pdf_bytes(CLEAN_LINE), "clean.pdf")
        assert resp.status_code == 200
        doc = fitz.open(stream=resp.content, filetype="pdf")
        full_text = "".join(page.get_text() for page in doc)
        assert CLEAN_LINE in full_text


class TestRedactFileDownloadDOCX:
    def test_returns_docx_content_type(self):
        resp = _post_download(_make_docx_bytes(SECRET_LINE), "report.docx")
        assert resp.status_code == 200
        assert "wordprocessingml" in resp.headers["content-type"]

    def test_filename_has_redacted_suffix(self):
        resp = _post_download(_make_docx_bytes(SECRET_LINE), "report.docx")
        assert "report_redacted.docx" in resp.headers["content-disposition"]

    def test_secret_replaced_in_returned_docx(self):
        from docx import Document
        resp = _post_download(_make_docx_bytes(SECRET_LINE), "secret.docx")
        assert resp.status_code == 200
        doc = Document(io.BytesIO(resp.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "sk-testsecretvalue1234567890abcdef" not in full_text
        assert "<SECRET>" in full_text

    def test_clean_docx_passes_through(self):
        from docx import Document
        resp = _post_download(_make_docx_bytes(CLEAN_LINE), "clean.docx")
        assert resp.status_code == 200
        doc = Document(io.BytesIO(resp.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert CLEAN_LINE in full_text

    def test_download_unsupported_type_returns_422(self):
        resp = _post_download(b"secret content", "data.txt")
        assert resp.status_code == 422
