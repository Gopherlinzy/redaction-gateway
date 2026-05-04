"""
文件文本提取与写回脱敏：PDF（PyMuPDF）和 DOCX（python-docx）。

公开 API：
  extract_text(path, suffix)         → str          纯文本提取
  redact_pdf_bytes(path, values)     → bytes         返回脱敏后的 PDF
  redact_docx_bytes(path, values)    → bytes         返回脱敏后的 DOCX
  UnsupportedFileTypeError           (ValueError)

所有库都是懒加载 —— ImportError 转为 UnsupportedFileTypeError。
"""

import io
from pathlib import Path


class UnsupportedFileTypeError(ValueError):
    pass


# ── 文本提取 ───────────────────────────────────────────────────────────────────

def _extract_pdf(path: str) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise UnsupportedFileTypeError("PyMuPDF not installed; cannot extract PDF") from exc

    doc = fitz.open(path)
    pages: list[str] = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\f\n".join(pages)


def _extract_docx(path: str) -> str:
    """
    只遍历 body 顶层元素，避免 doc.paragraphs 把表格内的段落计两次。
    遇到顶层 <w:p> 直接读；遇到顶层 <w:tbl> 递归读单元格段落。
    """
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise UnsupportedFileTypeError("python-docx not installed; cannot extract DOCX") from exc

    doc = Document(path)
    parts: list[str] = []

    def _collect_table(tbl: "Table") -> None:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t:
                        parts.append(t)
                for nested in cell.tables:
                    _collect_table(nested)

    for child in doc.element.body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local == "p":
            t = Paragraph(child, doc).text.strip()
            if t:
                parts.append(t)
        elif local == "tbl":
            _collect_table(Table(child, doc))

    return "\n".join(parts)


_EXTRACTORS = {
    ".pdf":  _extract_pdf,
    ".docx": _extract_docx,
}


def extract_text(path: str, suffix: str | None = None) -> str:
    ext = (suffix or Path(path).suffix).lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'; supported: {', '.join(_EXTRACTORS)}"
        )
    return extractor(path)


# ── 写回脱敏 ───────────────────────────────────────────────────────────────────

def _pdf_page_chars(page: "fitz.Page") -> list[tuple[str, "fitz.Rect"]]:
    """返回页面所有字符及其 bbox（rawdict），包含换行符占位以对齐 get_text() 输出。"""
    import fitz
    chars: list[tuple[str, fitz.Rect]] = []
    for block in page.get_text("rawdict")["blocks"]:
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                for ch in span.get("chars", []):
                    c = ch.get("c", "")
                    if c:
                        chars.append((c, fitz.Rect(ch["bbox"])))
            # get_text() 每行末尾加换行，rawdict 对齐需同步
            chars.append(("\n", fitz.Rect(0, 0, 0, 0)))
    return chars


def _find_value_rects(page: "fitz.Page", value: str) -> list["fitz.Rect"]:
    """
    先用 search_for 快速定位；失败时用字符级 bbox 回退（覆盖中文 CJK 字体场景）。
    两者都用 page.get_text() 同源文本，确保 find() 偏移量对齐。
    """
    import fitz

    rects = page.search_for(value)
    if rects:
        return rects

    # 回退：从 rawdict 取字符 bbox，用 get_text() 同源文本定位偏移
    chars = _pdf_page_chars(page)
    if not chars:
        return []

    # 用与提取时相同的 get_text() 文本做 find，确保偏移一致
    page_text = page.get_text()
    chars_filtered = [(c, r) for c, r in chars if c != "\n"]

    # 在 page_text 里找到所有匹配，映射到 chars_filtered（跳过换行符）
    # 重建不含换行的索引映射
    no_newline_text = page_text.replace("\n", "")
    value_no_newline = value.replace("\n", "")

    found: list[fitz.Rect] = []
    start = 0
    while True:
        idx = no_newline_text.find(value_no_newline, start)
        if idx == -1:
            break
        span_rects = [r for _, r in chars_filtered[idx: idx + len(value_no_newline)]
                      if r.width > 0 or r.height > 0]
        if span_rects:
            combined = span_rects[0]
            for r in span_rects[1:]:
                combined |= r
            found.append(combined)
        start = idx + 1
    return found


def redact_pdf_bytes(path: str, secret_values: list[str]) -> bytes:
    """
    用 PyMuPDF redaction API 物理删除敏感文字并覆盖黑块，返回新 PDF bytes。
    对 search_for 无法定位的 CJK 字体文字，回退到字符级 bbox 定位。
    """
    try:
        import fitz
    except ImportError as exc:
        raise UnsupportedFileTypeError("PyMuPDF not installed; cannot redact PDF") from exc

    doc = fitz.open(path)
    for page in doc:
        for value in secret_values:
            if not value:
                continue
            for rect in _find_value_rects(page, value):
                page.add_redact_annot(rect, fill=(0, 0, 0))
        page.apply_redactions()
    out = doc.tobytes(garbage=4, deflate=True)
    doc.close()
    return out


def redact_docx_bytes(path: str, secret_values: list[str]) -> bytes:
    """
    在段落/表格中替换密钥值，返回新 DOCX bytes。
    两步法：
      1. per-run 替换（保留各 run 的字体/加粗等格式）
      2. 若密钥横跨多个 run（Word 自动分拆），回退到段落级替换
         （第一个 run 承接所有文字，其余 run 清空 —— 丢失跨 run 的局部格式，
          但确保密钥不会漏过）
    """
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise UnsupportedFileTypeError("python-docx not installed; cannot redact DOCX") from exc

    doc = Document(path)
    values = [v for v in secret_values if v]

    def _redact_para(para: "Paragraph") -> None:
        for v in values:
            # pass 1: per-run（保格式）
            for run in para.runs:
                if v in run.text:
                    run.text = run.text.replace(v, "<SECRET>")
            # pass 2: 跨 run 兜底
            if v in para.text:
                replaced = para.text.replace(v, "<SECRET>")
                runs = para.runs
                if runs:
                    runs[0].text = replaced
                    for run in runs[1:]:
                        run.text = ""

    def _redact_table(tbl: "Table") -> None:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _redact_para(para)
                for nested in cell.tables:
                    _redact_table(nested)

    for child in doc.element.body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local == "p":
            _redact_para(Paragraph(child, doc))
        elif local == "tbl":
            _redact_table(Table(child, doc))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
